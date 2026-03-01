"""
DOCX file parser for Microsoft Word documents.
Uses python-docx library for text extraction.
"""

import os
from typing import Tuple
from .base_parser import BaseParser


class DocxParser(BaseParser):
    """Parser for Microsoft Word (.docx) files"""

    def __init__(self):
        """Initialize DOCX parser"""
        self.supported_extensions = {'.docx'}

    def parse(self, file_path: str) -> dict:
        """
        Parse DOCX file and extract text content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary with parsing results
        """
        # Validate file first
        is_valid, message = self.validate(file_path)
        if not is_valid:
            return {
                'success': False,
                'content': '',
                'metadata': {'file_path': file_path, 'error': message},
                'error': message
            }

        try:
            # Import docx here to avoid dependency issues
            from docx import Document

            # Load and parse document
            doc = Document(file_path)

            # Extract text from all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    paragraphs.append(paragraph.text)

            content = '\n'.join(paragraphs)

            # Get metadata
            metadata = self.get_metadata(file_path)

            return {
                'success': True,
                'content': content,
                'metadata': metadata,
                'error': None
            }

        except ImportError:
            return {
                'success': False,
                'content': '',
                'metadata': {'file_path': file_path},
                'error': "python-docx library not installed. Install with: pip install python-docx"
            }
        except Exception as e:
            return {
                'success': False,
                'content': '',
                'metadata': {'file_path': file_path},
                'error': f"Error parsing DOCX file: {str(e)}"
            }

    def get_metadata(self, file_path: str) -> dict:
        """
        Get metadata about the DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary with file metadata
        """
        metadata = self._get_basic_metadata(file_path)

        try:
            from docx import Document

            doc = Document(file_path)

            # Count paragraphs
            paragraph_count = len(doc.paragraphs)

            # Count non-empty paragraphs
            non_empty_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())

            # Get document properties if available
            core_properties = {}
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                core_properties = {
                    'title': props.title,
                    'author': props.author,
                    'created': str(props.created) if props.created else None,
                    'modified': str(props.modified) if props.modified else None,
                    'last_modified_by': props.last_modified_by,
                    'revision': props.revision,
                    'category': props.category,
                    'keywords': props.keywords,
                    'comments': props.comments,
                }

            metadata.update({
                'paragraph_count': paragraph_count,
                'non_empty_paragraph_count': non_empty_paragraphs,
                'document_properties': core_properties,
                'has_tables': len(doc.tables) > 0 if hasattr(doc, 'tables') else False,
                'table_count': len(doc.tables) if hasattr(doc, 'tables') else 0,
            })

        except ImportError:
            metadata['error'] = "python-docx library not installed"
        except Exception as e:
            metadata['error'] = f"Error reading DOCX metadata: {str(e)}"

        return metadata

    def validate(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate if file is a valid DOCX file.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (is_valid: bool, message: str)
        """
        # Check if file exists and is readable
        file_valid, file_message = self._validate_file_exists(file_path)
        if not file_valid:
            return False, file_message

        # Check file extension
        _, ext = os.path.splitext(file_path)
        if ext.lower() not in self.supported_extensions:
            return False, f"File extension {ext} is not supported. Expected .docx"

        # Check file signature (optional - could check for ZIP header)
        # Only check if file is large enough
        try:
            if os.path.getsize(file_path) >= 4:
                with open(file_path, 'rb') as f:
                    header = f.read(4)
                    # DOCX files are ZIP archives starting with PK\x03\x04
                    if header != b'PK\x03\x04':
                        return False, "File does not appear to be a valid DOCX (missing ZIP header)"
        except (IOError, OSError):
            pass  # Skip signature check if we can't read file

        return True, "Valid DOCX file"